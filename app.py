from flask import Flask, render_template, request, send_file, flash, redirect, url_for
import difflib
import os
import fitz  # PyMuPDF
import docx
from reportlab.pdfgen import canvas
from io import BytesIO
import magic

app = Flask(__name__)
app.secret_key = b'_5#y2Lzcgsfga"FzzzzQ8xxx\n\xdc]/'

# --- Conversion Functions ---

def docx_to_text(file_stream):
    doc = docx.Document(file_stream)
    return '\n'.join(para.text for para in doc.paragraphs)

def pdf_to_text(file_stream):
    doc = fitz.open(stream=file_stream.read(), filetype="pdf")
    text = ''
    for page in doc:
        text += page.get_text()
    return text

def text_to_docx(text):
    doc = docx.Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream

def text_to_pdf(text):
    output_stream = BytesIO()
    c = canvas.Canvas(output_stream)
    y = 800
    for line in text.splitlines():
        if y < 50:
            c.showPage()
            y = 800
        c.drawString(50, y, line)
        y -= 15
    c.save()
    output_stream.seek(0)
    return output_stream

# --- Routes ---

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/compare', methods=['POST'])
def compare():
    file1 = request.files['file1']
    file2 = request.files['file2']

    if not file1 or not file2:
        flash('Both files are required!')
        return redirect(url_for('index'))

    ext1 = os.path.splitext(file1.filename)[1].lower()
    ext2 = os.path.splitext(file2.filename)[1].lower()

    # Validate file types using MIME type
    mime = magic.Magic(mime=True)
    file1_content = file1.read()
    file1_mime = mime.from_buffer(file1_content)
    file1.seek(0)

    file2_content = file2.read()
    file2_mime = mime.from_buffer(file2_content)
    file2.seek(0)

    # Convert file1 to text
    if ext1 == '.docx':
        if file1_mime != 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            flash('File 1 is not a valid DOCX file!')
            return redirect(url_for('index'))
        text1 = docx_to_text(file1)
    elif ext1 == '.pdf':
        if file1_mime != 'application/pdf':
            flash('File 1 is not a valid PDF file!')
            return redirect(url_for('index'))
        text1 = pdf_to_text(file1)
    elif ext1 == '.txt':
        if not file1_mime.startswith('text/'):
            flash('File 1 is not a valid text file!')
            return redirect(url_for('index'))
        text1 = file1.read().decode('utf-8')
    else:
        flash('Unsupported file type for File 1')
        return redirect(url_for('index'))

    # Convert file2 to text
    if ext2 == '.docx':
        if file2_mime != 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            flash('File 2 is not a valid DOCX file!')
            return redirect(url_for('index'))
        text2 = docx_to_text(file2)
    elif ext2 == '.pdf':
        if file2_mime != 'application/pdf':
            flash('File 2 is not a valid PDF file!')
            return redirect(url_for('index'))
        text2 = pdf_to_text(file2)
    elif ext2 == '.txt':
        if not file2_mime.startswith('text/'):
            flash('File 2 is not a valid text file!')
            return redirect(url_for('index'))
        text2 = file2.read().decode('utf-8')
    else:
        flash('Unsupported file type for File 2')
        return redirect(url_for('index'))

    # Diff Comparison
    differ = difflib.HtmlDiff(wrapcolumn=80)
    diff_html = differ.make_table(text1.splitlines(), text2.splitlines(), fromdesc=file1.filename, todesc=file2.filename, context=True, numlines=3)

    # Store diff text for exporting
    global last_diff_text
    last_diff_text = difflib.unified_diff(text1.splitlines(), text2.splitlines(), lineterm='')
    last_diff_text = '\n'.join(list(last_diff_text))

    return render_template('result.html', diff_html=diff_html)

@app.route('/export/<fmt>')
def export(fmt):
    if 'last_diff_text' not in globals() or not last_diff_text:
        return 'No comparison done yet.'

    if fmt == 'txt':
        return send_file(BytesIO(last_diff_text.encode('utf-8')), download_name='diff.txt', as_attachment=True)
    elif fmt == 'docx':
        return send_file(text_to_docx(last_diff_text), download_name='diff.docx', as_attachment=True)
    elif fmt == 'pdf':
        return send_file(text_to_pdf(last_diff_text), download_name='diff.pdf', as_attachment=True)
    else:
        return 'Unsupported export format.'

if __name__ == '__main__':
    app.run(debug=True, port=8000)
